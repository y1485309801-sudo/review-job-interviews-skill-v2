#!/usr/bin/env python3
"""Build a polished interview review DOCX from structured JSON."""

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "203A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "3E6B78"
INK = "24313A"
MUTED = "66737D"
LIGHT_BLUE = "EAF1F7"
PALE_BLUE = "F4F8FB"
LIGHT_GRAY = "F2F4F7"
GREEN = "2F6B55"
PALE_GREEN = "EAF4EF"
AMBER = "8A6400"
PALE_AMBER = "FFF7E3"
RED = "9B3A3A"
PALE_RED = "FCEEEE"
WHITE = "FFFFFF"


LABELS = {
    "zh": {
        "report_title": "{company}面试复盘",
        "kicker": "INTERVIEW REVIEW",
        "based_on": "基于候选人简历、岗位描述与面试录音转写整理",
        "candidate": "候选人",
        "target_role": "目标岗位",
        "sources": "分析依据",
        "usage": "使用方式",
        "made_on": "制作日期",
        "section_summary": "一、核心结论",
        "verdict": "一句话结论",
        "strengths": "最值得保留的优势",
        "gaps": "必须修复的失分点",
        "match": "岗位要求与候选人证据映射",
        "requirement": "岗位要求",
        "evidence": "候选人证据",
        "fit": "匹配度",
        "strategy": "面试策略",
        "fact_boundary": "事实边界",
        "section_questions": "二、录音提炼问题与优化答案",
        "questions_note": "问题按面试顺序整理。蓝色框中的内容可直接练习。",
        "intent": "考察意图｜",
        "diagnosis": "现场复盘｜",
        "risk": "现场风险",
        "answer": "优化后的直接回答",
        "basis": "事实依据｜",
        "project_evidence": "本回答使用的项目档案证据",
        "newly_confirmed": "本次新确认",
        "rewrite_note": "重写说明｜",
        "followups": "可能追问",
        "followup_answer": "答：",
        "section_likely": "三、下一轮高概率问题",
        "likely_note": "以下问题来自岗位要求中尚未充分验证或现场暴露出的能力缺口。",
        "section_strategy": "四、表达与面试策略",
        "answer_structures": "统一回答结构",
        "risk_phrases": "需要替换的高风险表达",
        "avoid": "避免这样说",
        "replace": "改成这样说",
        "speaking_tips": "口头表达修正",
        "section_quick": "五、面试前速记卡",
        "positioning": "个人定位",
        "core_evidence": "核心证据",
        "boundaries": "经验边界",
        "reverse_questions": "反问面试官",
        "practice": "练习清单",
        "section_receipt": "六、本次项目档案变化回执",
        "receipt_note": "档案位置",
        "receipt_summary": "变更统计",
        "receipt_warning_title": "未确认内容",
        "receipt_pending": "仍有未确认内容；这些内容未进入正式档案或优化答案。",
        "update_id": "编号",
        "project": "项目",
        "change": "变化",
        "decision": "处理结果",
        "confirmed_content": "内容与来源",
        "affected_questions": "自动重写的问题",
        "source_prefix": "来源：",
        "footer": "面试复盘与回答手册",
        "page": "第 ",
        "page_suffix": " 页",
    },
    "en": {
        "report_title": "{company} Interview Review",
        "kicker": "INTERVIEW REVIEW",
        "based_on": "Prepared from the candidate resume, job description, and interview transcript",
        "candidate": "Candidate",
        "target_role": "Target role",
        "sources": "Sources",
        "usage": "How to use",
        "made_on": "Prepared",
        "section_summary": "1. Executive Summary",
        "verdict": "Bottom line",
        "strengths": "Strengths to keep",
        "gaps": "Priority gaps",
        "match": "Role-to-evidence mapping",
        "requirement": "Requirement",
        "evidence": "Candidate evidence",
        "fit": "Fit",
        "strategy": "Interview strategy",
        "fact_boundary": "Evidence boundary",
        "section_questions": "2. Transcript Questions and Improved Answers",
        "questions_note": "Questions follow the interview sequence. Blue callouts are ready to rehearse.",
        "intent": "Intent | ",
        "diagnosis": "Review | ",
        "risk": "High-risk statement",
        "answer": "Improved direct answer",
        "basis": "Evidence | ",
        "project_evidence": "Confirmed project evidence used",
        "newly_confirmed": "Newly confirmed",
        "rewrite_note": "Rewrite note | ",
        "followups": "Likely follow-ups",
        "followup_answer": "Answer: ",
        "section_likely": "3. Likely Next-Round Questions",
        "likely_note": "These questions cover JD requirements not fully tested or gaps revealed in the interview.",
        "section_strategy": "4. Communication and Interview Strategy",
        "answer_structures": "Answer structures",
        "risk_phrases": "Phrases to replace",
        "avoid": "Avoid",
        "replace": "Use instead",
        "speaking_tips": "Speaking improvements",
        "section_quick": "5. Pre-Interview Quick Card",
        "positioning": "Positioning",
        "core_evidence": "Core evidence",
        "boundaries": "Experience boundaries",
        "reverse_questions": "Questions for the interviewer",
        "practice": "Practice checklist",
        "section_receipt": "6. Project Profile Change Receipt",
        "receipt_note": "Profile location",
        "receipt_summary": "Change summary",
        "receipt_warning_title": "Unconfirmed items",
        "receipt_pending": "Some items remain unconfirmed; they were not added to the profile or optimized answers.",
        "update_id": "ID",
        "project": "Project",
        "change": "Change",
        "decision": "Decision",
        "confirmed_content": "Content and source",
        "affected_questions": "Questions automatically rewritten",
        "source_prefix": "Source: ",
        "footer": "Interview review and answer guide",
        "page": "Page ",
        "page_suffix": "",
    },
}

ENUM_LABELS = {
    "zh": {
        "total": "共计",
        "confirmed": "已确认",
        "edited": "修改后确认",
        "rejected": "已拒绝",
        "pending": "待确认",
        "applied": "已入档",
        "preview_only": "仅预览",
        "new": "新增",
        "enrich": "补充",
        "reinforce": "再次确认",
        "conflict": "冲突处理",
        "add_project": "新增项目",
        "add_fact": "新增事实",
        "reinforce_fact": "补充来源",
        "replace_fact": "替换事实",
        "actions": "个人行动",
        "results": "结果",
        "metrics": "指标",
        "tech_stack": "技术工具",
        "deliverables": "交付物",
    },
    "en": {},
}


def localize_enum(language, value):
    return ENUM_LABELS.get(language, {}).get(str(value), str(value))


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    latin = "Calibri"
    east_asia = "Hiragino Sans GB"
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_para_left_border(paragraph, color=BLUE, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_field(paragraph, label, suffix):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    before = paragraph.add_run(label)
    set_run_font(before, size=9, color=MUTED)
    field_run = paragraph.add_run()
    start = OxmlElement("w:fldChar")
    start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.extend([start, instruction, end])
    after = paragraph.add_run(suffix)
    set_run_font(after, size=9, color=MUTED)


def add_numbering(doc, marker="•"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", marker), ("w:lvlJc", "left")):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    return section


def add_bullet(doc, num_id, text, color=INK):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph._p.get_or_add_pPr().append(num_pr)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=11, color=color)
    return paragraph


def add_label(doc, label, text, label_color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    first = paragraph.add_run(label)
    set_run_font(first, size=10.5, color=label_color, bold=True)
    second = paragraph.add_run(str(text))
    set_run_font(second, size=10.5, color=INK)
    return paragraph


def add_callout(doc, title, text, fill=PALE_BLUE, border=BLUE, title_color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.28
    set_para_shading(paragraph, fill)
    set_para_left_border(paragraph, border)
    heading = paragraph.add_run(str(title) + "\n")
    set_run_font(heading, size=10.5, color=title_color, bold=True)
    body = paragraph.add_run(str(text))
    set_run_font(body, size=11, color=INK)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(str(text))
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return paragraph


def add_section_heading(doc, text):
    paragraph = add_heading(doc, text, 1)
    paragraph.paragraph_format.page_break_before = True
    return paragraph


def add_data_table(doc, headers, rows, widths, centered_columns=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = str(text)
        set_cell_fill(cell, NAVY)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            if row_index % 2 == 0:
                set_cell_fill(cells[index], PALE_BLUE)
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in centered_columns else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.3, color=INK)
    return table


def add_question(doc, labels, num_id, index, item):
    question_id = item.get("question_id") or f"Q{index:02d}"
    add_heading(doc, f"{question_id}  {item.get('question', '')}", 2)
    add_label(doc, labels["intent"], item.get("intent", ""))
    if item.get("risk"):
        add_callout(doc, labels["risk"], item["risk"], PALE_RED, RED, RED)
    if item.get("diagnosis"):
        add_label(doc, labels["diagnosis"], item["diagnosis"], AMBER)
    add_callout(doc, labels["answer"], item.get("answer", ""))
    if item.get("evidence"):
        add_label(doc, labels["basis"], item["evidence"], GREEN)
    project_evidence = item.get("project_evidence") or []
    if project_evidence:
        add_heading(doc, labels["project_evidence"], 3)
        for project in project_evidence:
            fact_parts = []
            for fact in project.get("facts_used") or []:
                marker = (
                    f" [{labels['newly_confirmed']}]"
                    if fact.get("newly_confirmed")
                    else ""
                )
                fact_id = fact.get("fact_id", "")
                prefix = f"{fact_id}: " if fact_id else ""
                fact_parts.append(f"{prefix}{fact.get('fact', '')}{marker}")
            project_name = project.get("project_name") or project.get("project_id", "")
            text = f"{project_name}｜" + "；".join(fact_parts)
            add_bullet(doc, num_id, text)
    if item.get("rewrite_note"):
        add_label(doc, labels["rewrite_note"], item["rewrite_note"], TEAL)
    followups = item.get("followups") or []
    if followups:
        add_heading(doc, labels["followups"], 3)
        for followup in followups:
            text = f"{followup.get('question', '')}  {labels['followup_answer']}{followup.get('answer', '')}"
            add_bullet(doc, num_id, text)


def build_report(data, output):
    meta = data.get("meta") or {}
    summary = data.get("summary") or {}
    language = str(meta.get("language", "zh")).lower()
    language = "en" if language.startswith("en") else "zh"
    labels = LABELS[language]
    company = meta.get("company") or ("Target Company" if language == "en" else "目标公司")
    role = meta.get("role") or ("Target Role" if language == "en" else "目标岗位")
    candidate = meta.get("candidate") or ("Candidate" if language == "en" else "候选人")
    prepared = meta.get("date") or date.today().isoformat()

    doc = Document()
    section = configure_document(doc)
    bullet_id = add_numbering(doc)

    header = section.header
    header_p = header.paragraphs[0]
    run = header_p.add_run(f"{company}｜{role}")
    set_run_font(run, size=9, color=MUTED, bold=True)
    footer = section.footer
    footer_run = footer.paragraphs[0].add_run(labels["footer"])
    set_run_font(footer_run, size=9, color=MUTED)
    add_page_field(footer.add_paragraph(), labels["page"], labels["page_suffix"])

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run(labels["kicker"]), size=11, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    title_text = meta.get("title") or labels["report_title"].format(company=company)
    set_run_font(p.add_run(title_text), size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    subtitle = f"{role}｜{meta.get('role_meta', '')}".rstrip("｜")
    set_run_font(p.add_run(subtitle), size=15, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.space_after = Pt(30)
    lead = meta.get("source_note") or labels["based_on"]
    set_run_font(p.add_run(lead), size=11.5, color=MUTED)

    meta_rows = [
        (labels["candidate"], candidate),
        (labels["target_role"], role),
        (labels["sources"], meta.get("source_note") or labels["based_on"]),
        (labels["usage"], meta.get("usage_note") or labels["questions_note"]),
    ]
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2300, 7060])
    for row_index, (key, value) in enumerate(meta_rows):
        table.cell(row_index, 0).text = str(key)
        table.cell(row_index, 1).text = str(value)
        set_cell_fill(table.cell(row_index, 0), LIGHT_BLUE)
        for col_index, cell in enumerate(table.rows[row_index].cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for cell_run in paragraph.runs:
                    set_run_font(cell_run, size=10.5, color=DARK_BLUE if col_index == 0 else INK, bold=col_index == 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    set_run_font(p.add_run(f"{labels['made_on']}：{prepared}"), size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, labels["section_summary"], 1)
    if summary.get("verdict"):
        add_callout(doc, labels["verdict"], summary["verdict"], PALE_GREEN, GREEN, GREEN)
    if summary.get("score_line"):
        add_label(doc, "", summary["score_line"])
    if summary.get("strengths"):
        add_heading(doc, labels["strengths"], 2)
        for item in summary["strengths"]:
            add_bullet(doc, bullet_id, item)
    if summary.get("gaps"):
        add_heading(doc, labels["gaps"], 2)
        for item in summary["gaps"]:
            add_bullet(doc, bullet_id, item)
    match_rows = summary.get("match_rows") or []
    if match_rows:
        add_heading(doc, labels["match"], 2)
        rows = [[row.get("requirement", ""), row.get("evidence", ""), row.get("fit", ""), row.get("strategy", "")] for row in match_rows]
        add_data_table(
            doc,
            [labels["requirement"], labels["evidence"], labels["fit"], labels["strategy"]],
            rows,
            [1950, 3650, 1200, 2560],
            centered_columns=(2,),
        )
    if summary.get("fact_boundary"):
        add_heading(doc, labels["fact_boundary"], 2)
        add_callout(doc, labels["fact_boundary"], summary["fact_boundary"], PALE_AMBER, AMBER, AMBER)

    questions = data.get("questions") or []
    if questions:
        add_section_heading(doc, labels["section_questions"])
        p = doc.add_paragraph()
        set_run_font(p.add_run(labels["questions_note"]), size=10.5, color=MUTED, italic=True)
        for index, item in enumerate(questions, start=1):
            add_question(doc, labels, bullet_id, index, item)

    likely = data.get("likely_questions") or []
    if likely:
        add_section_heading(doc, labels["section_likely"])
        add_callout(doc, labels["section_likely"], labels["likely_note"], PALE_AMBER, AMBER, AMBER)
        start = len(questions) + 1
        for offset, item in enumerate(likely):
            add_question(doc, labels, bullet_id, start + offset, item)

    strategy = data.get("strategy") or {}
    if any(strategy.get(key) for key in ("answer_structures", "risky_phrases", "speaking_tips")):
        add_section_heading(doc, labels["section_strategy"])
        if strategy.get("answer_structures"):
            add_heading(doc, labels["answer_structures"], 2)
            for item in strategy["answer_structures"]:
                add_callout(doc, item.get("title", ""), item.get("body", ""))
        if strategy.get("risky_phrases"):
            add_heading(doc, labels["risk_phrases"], 2)
            rows = [[item.get("avoid", ""), item.get("replace", "")] for item in strategy["risky_phrases"]]
            table = add_data_table(doc, [labels["avoid"], labels["replace"]], rows, [4450, 4910])
            for row in table.rows[1:]:
                for run in row.cells[0].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor.from_string(RED)
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor.from_string(GREEN)
        if strategy.get("speaking_tips"):
            add_heading(doc, labels["speaking_tips"], 2)
            for item in strategy["speaking_tips"]:
                add_bullet(doc, bullet_id, item)

    quick = data.get("quick_card") or {}
    if any(quick.get(key) for key in ("positioning", "core_evidence", "keyword_groups", "boundaries", "reverse_questions", "practice")):
        add_section_heading(doc, labels["section_quick"])
        if quick.get("positioning"):
            add_callout(doc, labels["positioning"], quick["positioning"], PALE_GREEN, GREEN, GREEN)
        if quick.get("core_evidence"):
            add_heading(doc, labels["core_evidence"], 2)
            for item in quick["core_evidence"]:
                add_bullet(doc, bullet_id, item)
        for group in quick.get("keyword_groups") or []:
            add_heading(doc, group.get("title", ""), 2)
            for item in group.get("items") or []:
                add_bullet(doc, bullet_id, item)
        if quick.get("boundaries"):
            add_heading(doc, labels["boundaries"], 2)
            add_callout(doc, labels["boundaries"], quick["boundaries"], PALE_AMBER, AMBER, AMBER)
        if quick.get("reverse_questions"):
            add_heading(doc, labels["reverse_questions"], 2)
            for item in quick["reverse_questions"]:
                add_bullet(doc, bullet_id, item)
        if quick.get("practice"):
            add_heading(doc, labels["practice"], 2)
            for item in quick["practice"]:
                add_bullet(doc, bullet_id, item)

    receipt = data.get("profile_update_receipt") or {}
    if receipt:
        add_heading(doc, labels["section_receipt"], 1)
        profile_note = receipt.get("profile_path_note") or "Private local profile"
        add_callout(
            doc,
            labels["receipt_note"],
            profile_note,
            PALE_GREEN,
            GREEN,
            GREEN,
        )
        counts = receipt.get("counts") or {}
        if counts:
            count_order = ("total", "confirmed", "edited", "rejected", "pending", "applied")
            count_text = " | ".join(
                f"{localize_enum(language, key)}: {counts.get(key, 0)}"
                for key in count_order
            )
            add_label(doc, labels["receipt_summary"] + "｜", count_text, TEAL)
        if counts.get("pending", 0):
            add_callout(
                doc,
                labels["receipt_warning_title"],
                labels["receipt_pending"],
                PALE_AMBER,
                AMBER,
                AMBER,
            )
        affected = receipt.get("affected_question_ids") or []
        if affected:
            add_label(
                doc,
                labels["affected_questions"] + "｜",
                "、".join(str(item) for item in affected),
                TEAL,
            )
        changes = receipt.get("changes") or []
        if changes:
            rows = []
            for change in changes:
                source = change.get("source") or {}
                source_text = source.get("source_label") or source.get("source_type", "")
                value_text = str(change.get("value", ""))
                if source_text:
                    value_text = (
                        f"{value_text}\n{labels['source_prefix']}{source_text}"
                    )
                change_text = " / ".join(
                    localize_enum(language, value)
                    for value in (
                        change.get("change_type", ""),
                        change.get("field", ""),
                    )
                    if value
                )
                decision_text = " / ".join(
                    localize_enum(language, value)
                    for value in (
                        change.get("decision", ""),
                        change.get("status", ""),
                    )
                    if value
                )
                rows.append(
                    [
                        change.get("update_id", ""),
                        change.get("project_name") or change.get("project_id", ""),
                        change_text,
                        decision_text,
                        value_text,
                    ]
                )
            add_data_table(
                doc,
                [
                    labels["update_id"],
                    labels["project"],
                    labels["change"],
                    labels["decision"],
                    labels["confirmed_content"],
                ],
                rows,
                [800, 1500, 1450, 1400, 4210],
                centered_columns=(0, 3),
            )

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    doc.core_properties.title = title_text
    doc.core_properties.subject = "Interview review generated from resume, job description, and transcript"
    doc.core_properties.author = "Codex"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="UTF-8 report JSON")
    parser.add_argument("--output", required=True, type=Path, help="Destination DOCX")
    args = parser.parse_args()
    with args.input.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise SystemExit("Report JSON must be an object")
    build_report(data, args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
